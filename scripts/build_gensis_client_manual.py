from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "client-manual"
SHOT_DIR = OUT_DIR / "screenshots"
OUT_DOCX = OUT_DIR / "GENSIS ERP Client Manual.docx"

BRAND = "GENSIS ERP"
INK = RGBColor(15, 61, 51)
BLUE = RGBColor(45, 111, 150)
MUTED = RGBColor(89, 104, 122)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, bold=None, color=None, size=None) -> None:
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if level == 1:
            set_run(run, bold=True, color=BLUE, size=16)
        elif level == 2:
            set_run(run, bold=True, color=BLUE, size=13)
        else:
            set_run(run, bold=True, color=INK, size=12)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT_FILL)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run(run, bold=True, color=INK)
    paragraph.add_run(f"\n{text}")
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_widths = [Inches(width) for width in widths]
    set_table_geometry(table, table_widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(header)
        set_run(run, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            set_cell_margins(cells[index])
    doc.add_paragraph()


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SHOT_DIR / filename
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.4))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(10)
    for run in caption_paragraph.runs:
        run.italic = True
        run.font.color.rgb = MUTED
        run.font.size = Pt(9)


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(72)
    logo = paragraph.add_run("G")
    set_run(logo, bold=True, color=INK, size=44)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{BRAND}\nClient Review Manual")
    set_run(run, bold=True, color=INK, size=28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Feature overview, workflow guide, screenshots, and client feedback checklist")
    set_run(run, color=MUTED, size=12)

    add_callout(
        doc,
        "Document Purpose",
        "This manual is intended for client review. It explains what is currently available in GENSIS ERP, how key workflows operate, and where the client should confirm preferences, exclusions, and future requirements."
    )
    meta = [
        ["Product", BRAND],
        ["Audience", "Client stakeholders, operations leads, finance/admin users, implementation reviewers"],
        ["Status", "Client review build"],
        ["Demo URL", "http://localhost:3000"],
        ["Demo login", "admin@eversafe-demo.test / Admin123!"]
    ]
    add_table(doc, ["Field", "Details"], meta, [1.6, 4.9])
    doc.add_page_break()


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    add_cover(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "GENSIS ERP is a single-company ERP workspace covering company administration, master data, sales and CRM, purchasing, inventory, accounting, HR, approvals, reporting, and integration foundations. The system is designed for an administrator-led setup where super admin users define company identity, branches, users, roles, and permissions before operational users begin daily work."
    )
    add_bullets(doc, [
        "One unified web interface with module navigation, branch selector, quick-create menu, notifications, dark mode, search, and logout controls.",
        "Super admin configuration for company profile, logo/icon, branches, departments, warehouses, cost centers, users, roles, and feature-level permissions.",
        "Operational workflows for master data, sales orders, invoices, purchasing, stock movement, accounting entries, HR requests, approval routing, and reports.",
        "Management reporting with dashboard, operational, financial, export, email scheduling, and Tally voucher export foundations."
    ])
    add_screenshot(doc, "01-login.png", "Figure 1: GENSIS ERP login page with branded entry point.")
    add_screenshot(doc, "02-dashboard.png", "Figure 2: GENSIS ERP management dashboard after login.")

    add_heading(doc, "2. End-to-End System Flow", 1)
    add_body(doc, "The recommended implementation flow starts with administration and setup, then progresses into operational data entry, finance posting, and reporting.")
    add_numbered(doc, [
        "Super admin signs in and validates company branding, address, currency, date format, fiscal year, invoice text, and terms.",
        "Admin creates branches, departments, warehouses, and cost centers that represent the client organisation.",
        "Admin defines roles with checkbox-selected feature permissions and creates users with the appropriate roles.",
        "Master-data users create or import customers, suppliers, items, categories, units, tax codes, currencies, payment terms, and warehouse bins.",
        "Sales and purchasing users create enquiries, quotations, sales orders, purchase requests, RFQs, purchase orders, deliveries, receipts, invoices, and returns.",
        "Inventory users post opening stock, transfers, counts, adjustments, and bulk Excel imports where applicable.",
        "Accounting users post source documents to the general ledger, manage journals, supplier payments, tax, bank reconciliation, budgets, opening balances, and year-end close.",
        "HR and approval users manage employees, leave, attendance, expenses, approval routing, notifications, checklists, holidays, and certification alerts.",
        "Managers review dashboards, operational reports, financial reports, exports, scheduled email deliveries, and integration status."
    ])

    add_heading(doc, "3. Module Coverage", 1)
    add_table(doc, ["Module", "Current Features", "Client Review Focus"], [
        ["Dashboard", "User, role, branch, department, warehouse, audit, failed login, and approval summary cards.", "Confirm KPI cards and any additional executive metrics."],
        ["Company", "Company identity, email, address, logo/icon, currency, timezone, date format, fiscal year, invoice footer, terms.", "Confirm official company details, invoice text, logo, and branding."],
        ["Organisation", "Branches, departments, warehouses, cost centers, active status tracking.", "Confirm whether hierarchy, branch codes, warehouse naming, and cost centers are sufficient."],
        ["Users and Roles", "Create users, assign roles, create custom roles, checkbox permissions, protected system roles.", "Confirm final role matrix and approval authority levels."],
        ["Master Data", "Customers, suppliers, items, supporting masters, attachments, saved views, CSV import/export, validation.", "Confirm required item/customer/supplier fields and import templates."],
        ["Sales and CRM", "Leads, opportunities, enquiries, quotations, revisions, sales orders, deliveries, invoices, receipts, credit notes, returns.", "Confirm quotation/invoice print format, taxes, discounts, and approval needs."],
        ["Purchasing", "Purchase requests, RFQs, supplier quotes, comparison, purchase orders, goods receipts, supplier invoices, returns.", "Confirm purchase approval rules, supplier invoice matching tolerances, and document formats."],
        ["Inventory", "Opening stock, stock ledger, transfers, counts, adjustments, valuation, Excel bulk import.", "Confirm barcode/bin needs, stock costing rules, and import formats."],
        ["Accounting", "Chart of accounts, periods, journals, GL posting, payments, trial balance, ageing, tax, statements, bank rec, budgets, recurring journals, year-end close.", "Confirm chart of accounts, tax rules, financial statement layouts, and close process."],
        ["HR and Approvals", "Employees, leave, attendance, expenses, approvals, notifications, checklists, holidays, certification alerts.", "Confirm HR fields, leave policies, approval stages, and escalation rules."],
        ["Reports and Integrations", "Dashboard, operational, financial reports, CSV/PDF exports, email schedule logs, Tally voucher export, Swagger API docs.", "Confirm report formats, schedule recipients, SMTP details, and Tally integration approach."]
    ], [1.35, 3.0, 2.15])

    add_heading(doc, "4. Super Admin Setup Manual", 1)
    add_heading(doc, "4.1 Company Branding and Basic Settings", 2)
    add_body(doc, "The Company screen is the primary setup area for basic customization. A super admin can update company name, email, phone, website, address, registration and tax numbers, base currency, timezone, date format, financial year start month, logo/icon, invoice footer, and terms.")
    add_screenshot(doc, "03-company-customization.png", "Figure 3: Company customization screen for basic identity and branding.")

    add_heading(doc, "4.2 Branches and Organisation Units", 2)
    add_body(doc, "The Organisation screen lets admin users add branches, departments, warehouses, and cost centers. Departments and warehouses can be tied to a parent branch where needed.")
    add_screenshot(doc, "04-organisation-setup.png", "Figure 4: Organisation setup screen for branches, departments, warehouses, and cost centers.")

    add_heading(doc, "4.3 Role and Permission Design", 2)
    add_body(doc, "Roles control what a user can see or edit. System roles are predefined and protected. Custom roles can be created with checkbox-selected feature permissions, such as sales view, sales edit, reports export, accounting edit, HR view, branch create, and other module-specific rights.")
    add_bullets(doc, [
        "Super admin or system administrator: full access through the global wildcard permission.",
        "System roles: protected defaults for common job functions.",
        "Custom roles: editable permission sets for client-specific access levels.",
        "Permission checkboxes: allow precise selection of the features available to each role."
    ])
    add_screenshot(doc, "05-role-permissions.png", "Figure 5: Custom role creation with filtered feature-permission checkboxes.")

    add_heading(doc, "4.4 User Creation and Role Assignment", 2)
    add_body(doc, "The Users screen allows super admin users to create new users with display name, email, username, temporary password, status, and one or more assigned roles. User access is derived from the permissions inside assigned roles.")
    add_screenshot(doc, "06-user-role-assignment.png", "Figure 6: User creation screen with role assignment checkboxes.")

    add_heading(doc, "5. Operational Workflow Manual", 1)
    add_heading(doc, "5.1 Sales, Invoices, Receipts, Credit Notes, and Returns", 2)
    add_body(doc, "Sales and CRM users can progress from lead and opportunity management into quotations, sales orders, deliveries, invoices, receipts, credit notes, and sales returns. Sales invoices can be posted to accounting through the source posting workflow.")
    add_screenshot(doc, "07-sales-finance.png", "Figure 7: Sales finance workspace with invoices, receipts, credit notes, and returns.")

    add_heading(doc, "5.2 Accounting and Financial Statements", 2)
    add_body(doc, "Accounting includes chart setup, financial periods, manual journals, source posting, supplier payments, bank reconciliation, tax summary, recurring journals, budgets, opening balances, year-end close, trial balance, ageing, balance sheet, profit and loss, and cash flow.")
    add_screenshot(doc, "08-accounting-statements.png", "Figure 8: Accounting statements view showing balance sheet, profit and loss, and cash flow.")

    add_heading(doc, "5.3 Inventory and Bulk Import", 2)
    add_body(doc, "Inventory users can review stock balances, stock ledger, inventory valuation, opening stock, transfers, counts, adjustments, and bulk import workflows for transfers, counts, and adjustments.")
    add_screenshot(doc, "11-inventory.png", "Figure 9: Inventory workspace and stock operations.")

    add_heading(doc, "5.4 Reporting and Integration Workflows", 2)
    add_body(doc, "Reports are available in dashboard, operational, financial, and integrations tabs. Reports can be exported as CSV or PDF, emailed on demand, scheduled for recurring delivery, and used as a foundation for Tally voucher export.")
    add_screenshot(doc, "09-financial-reports.png", "Figure 10: Financial reports workspace with trial balance, receivables, and payables.")
    add_screenshot(doc, "10-report-integrations.png", "Figure 11: Report integrations workspace with scheduled report delivery.")

    add_heading(doc, "6. Client Review Questions", 1)
    add_body(doc, "Use this section during client review meetings to confirm final requirements and identify anything the client does not want.")
    add_table(doc, ["Area", "Questions for Client", "Decision / Notes"], [
        ["Branding", "What final product name, company logo, colors, email, address, footer, and terms should be used?", ""],
        ["Organisation", "How many branches, departments, warehouses, and cost centers are needed at go-live?", ""],
        ["User Access", "Which roles are required? Which users need create/edit/export/admin rights?", ""],
        ["Master Data", "What customer, supplier, item, tax, unit, currency, and payment-term fields are mandatory?", ""],
        ["Sales", "What quotation, sales order, delivery, invoice, credit note, and receipt formats are required?", ""],
        ["Purchasing", "What purchase approval flow, RFQ process, PO format, and three-way matching tolerance are required?", ""],
        ["Inventory", "Is barcode, batch/serial tracking, bin management, reorder logic, or costing customization required?", ""],
        ["Accounting", "What chart of accounts, tax rules, fiscal periods, bank accounts, and financial statement formats are required?", ""],
        ["HR", "What employee fields, leave policies, attendance rules, approval stages, and escalation timings are required?", ""],
        ["Reports", "Which reports are must-have? Who receives scheduled reports, and in what format?", ""],
        ["Integrations", "Is SMTP available? Is Tally integration file export enough, or is direct delivery needed?", ""],
        ["Exclusions", "Which modules, screens, reports, or workflows should be hidden or removed?", ""]
    ], [1.25, 3.55, 1.7])

    add_heading(doc, "7. Current Decisions and Production Readiness", 1)
    add_callout(
        doc,
        "Important",
        "The current build is suitable for client review and workflow confirmation. Before production, the client should confirm final data migration, live email settings, backup/restore testing, hosting environment, security policies, and any legally required document formats."
    )
    add_table(doc, ["Topic", "Current Status", "Recommended Next Step"], [
        ["SMTP email", "Report scheduling and delivery logging are implemented. Local demo skips actual send when SMTP is not configured.", "Collect SMTP provider, sender address, credentials, and allowed recipients."],
        ["Tally", "Posted accounting journals can be exported as Tally-style vouchers.", "Confirm Tally version, import format, ledger mapping, and whether direct outbound delivery is required."],
        ["PDF layouts", "Basic PDF export exists for reports.", "Confirm branded invoice, PO, receipt, credit note, and financial report layouts."],
        ["Backup restore", "Listed as a remaining production test.", "Run backup/restore rehearsal before go-live."],
        ["Roles", "System roles and custom roles exist.", "Client should sign off final role-permission matrix."],
        ["Data migration", "CSV and Excel import foundations exist.", "Client should provide sample customer, supplier, item, stock, and opening-balance files."],
        ["Security", "Password hashing, cookies, permissions, audit logs, and MFA setup exist.", "Confirm password policy, MFA requirement, session timeout, and admin access process."]
    ], [1.4, 2.75, 2.35])

    add_heading(doc, "8. Acceptance Checklist", 1)
    add_bullets(doc, [
        "Client confirms GENSIS ERP product name and final logo.",
        "Client confirms company profile, address, email, tax details, invoice footer, and terms.",
        "Client confirms branches, departments, warehouses, and cost centers.",
        "Client confirms user list, roles, and permission checkboxes.",
        "Client confirms required sales, purchasing, inventory, accounting, HR, and reports workflows.",
        "Client confirms what should be removed or hidden from go-live.",
        "Client confirms required imports, exports, PDF layouts, email settings, and integrations.",
        "Client approves production-readiness tasks: deployment, security, backup restore, migration rehearsal, and UAT signoff."
    ])

    add_heading(doc, "Appendix A. Verified Build Notes", 1)
    add_body(doc, "Recent verification covered typechecking, automated tests, production builds, API smoke tests, and browser smoke tests for finance, accounting, reporting, top-bar controls, company customization, organisation setup, role permission checkboxes, and user role assignment.")
    add_table(doc, ["Verified Area", "Result"], [
        ["Finance workflow", "Sales order, invoice, receipt, credit note, sales return, and GL posting verified."],
        ["Accounting reports", "Trial balance, ageing, tax summary, financial statements, CSV/PDF report exports verified."],
        ["Admin customization", "Company branding/contact update, branch creation, custom role create/update, user create with role verified."],
        ["Browser screens", "Company, Organisation, Roles, Users, Reports, Sales Finance, Accounting, and Dashboard screens verified."],
        ["Known local limitation", "Actual report email sending requires SMTP configuration; local demo records SKIPPED delivery status when SMTP is absent."]
    ], [2.1, 4.4])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{BRAND} Client Review Manual")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
